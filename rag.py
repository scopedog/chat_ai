import os
import sys
import time
import requests
import re
import random
import string
import copy
import xml.etree.ElementTree as ET
import backoff
import fleep
import pandas as pd
import pdfplumber
from dotenv import load_dotenv
from loguru import logger
from pathlib import Path
from typing import Optional
from functools import partial
from bs4 import BeautifulSoup

# Set env
import os
os.environ["USER_AGENT"] = "shiningheresy"

from langchain.text_splitter import CharacterTextSplitter
from langchain_text_splitters import RecursiveCharacterTextSplitter
#from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import UnstructuredURLLoader
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders import WebBaseLoader
from langchain_community.document_loaders import PyPDFLoader
from langchain_docling import DoclingLoader
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from langchain_community.document_loaders import UnstructuredODTLoader
from langchain_community.document_loaders import UnstructuredExcelLoader
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain_community.vectorstores.utils import filter_complex_metadata
from chromadb.api.client import SharedSystemClient
import chromadb
from langchain.schema import Document
from langchain_chroma import Chroma
from docx import Document as Doc
from odf.opendocument import load
from odf.table import Table, TableRow, TableCell
from odf.text import P
from my_web_loader import MyWebLoader

# Class CtxDocs
class CtxDocs:
    def __init__(
        self,
        docs: list[Document] = None,
        combined_ctx: str = ""
    ):
        # Initialize
        self.docs = docs
        self.combined_ctx = combined_ctx

# Load RAG data
def load_data(
              data: list[str] = None, # Files, URLs
              raw_data: list[str] = None, # Raw text, html, ...
              chunk_size = 1024,
              chunk_overlap = 120,
              max_docs: int = 0
              ) -> CtxDocs:
    # Initialize
    urls = []
    files = []
    docs = []
    combined_ctx = ""

    # Check each item in data
    if bool(data):
        for item in data:
            if item.startswith("http://") or item.startswith("https://"):
                urls.append(item)
            else:
                files.append(item)

    # Load files
    if len(files) > 0:
        ctx_docs = load_files(files=files,
                              chunk_size=chunk_size,
                              chunk_overlap=chunk_overlap)
        docs.extend(ctx_docs.docs)
        combined_ctx += ctx_docs.combined_ctx

    # Load websites
    if len(urls) > 0:
        try:
            ctx_docs = load_websites(urls=urls,
                                     chunk_size=chunk_size,
                                     chunk_overlap=chunk_overlap,
                                     max_links=20)
            docs.extend(ctx_docs.docs)
            combined_ctx += ctx_docs.combined_ctx
        except Exception as e:
            raise Exception(e)

    # Load raw data
    if bool(raw_data):
        docs.extend(load_raw_data(raw_data))

    return CtxDocs(docs=docs, combined_ctx=combined_ctx)

# Load RAG files
def load_files(files: list[str],
               chunk_size = 1024,
               chunk_overlap = 120):

    docs = []

    # Create text splitter
    #print("chunk_size: " + str(chunk_size))
    #print("chunk_overlap: " + str(chunk_overlap))
    text_splitter = RecursiveCharacterTextSplitter(
                        separators=[
                            ".\n\n",
                            "\n\n",
                            "\r\n\r\n",
                            "。",
                        ],
                        keep_separator='end',
                        chunk_size=chunk_size,
                        chunk_overlap=chunk_overlap,
                        add_start_index=True,
                        ) 

    # Load RAG files
    for rag_file in files:
        if rag_file is not None:
            load_rag_file(rag_file=rag_file, docs=docs)

    # docs can be empty
    if len(docs) == 0:
        # If empty, add dummy doc
        doc = Document(
                page_content="",
                metadata={"source": "rag.py",
                          "page": 1,
                          "author": "HN"}
             )
        docs.append(doc)

    # Now docs is not split yet
    docs_no_split = docs

    # Combine page_content of docs_no_split
    combined_ctx = ""
    for d in docs_no_split:
        combined_ctx += d.page_content + "\n"
    logger.debug("combined_ctx:\n" + combined_ctx)

    # Split docs into chunks
    docs = text_splitter.split_documents(docs)
    #print(docs)

    #return docs, combined_ctx
    return CtxDocs(docs=docs, combined_ctx=combined_ctx)
 
# Load RAG file
# This doesn't split text
def load_rag_file(rag_file: str, docs):
    file = Path(rag_file)
    if not file.is_file():
        logger.warning("Cannot open " + rag_file)
        return
 
    # Check file type and load
    try:
        f = open(rag_file, "rb")
    except FileNotFoundError as e:
        loggoer.error(e)
        return

    info = fleep.get(f.read(128))
    f.close()
    #print(info.extension)

    # PDF
    if info.extension_matches("pdf"):
        logger.info("RAG file " + rag_file + " was detected as a PDF file")

        # We first use PyPDFLoader
        # This is good for retrieving data that are not included in tables
        '''
        loader = PyPDFLoader(rag_file)
        #load_pages_and_append(loader, docs)
        for page in loader.lazy_load():
            docs.append(page)
            #print("####\n" + str(page))
        '''

        # Use pdfplumber
        load_pdf_and_append(rag_file, docs)
        #print(docs)
        '''
        '''

    # XLS, XLSX
    # Caution! Put this before DOC as xlsx is detected as DOC
    elif info.extension_matches("xls") or info.extension_matches("xlsx"):
        logger.info("RAG file " + rag_file + " was detected as an XLS/XLSX file")
        # UnstructuredExcelLoader
        # For Gemma, this seems to be better than panda
        loader = UnstructuredExcelLoader(rag_file, mode="elements")
        #loader = UnstructuredExcelLoader(rag_file)
        load_pages_and_append(loader, docs)
        #print(docs)

        '''
        # pandas -- seems better than UnstructuredExcelLoader
        df = pd.read_excel(rag_file)
        docs_ = [Document(page_content=df.to_string(), metadata={"source": rag_file})]
        #print("docs_len: " + str(len(docs)))
        for doc in docs_:
            docs.append(doc)
        '''
        
    # DOCX
    elif info.extension_matches("docx"):
        logger.info("RAG file " + rag_file + " was detected as a DOCX file")

        # Use UnstructuredWordDocumentLoader first
        loader = UnstructuredWordDocumentLoader(rag_file)
        load_pages_and_append(loader, docs)

        # Retrieve tables and append
        tables = extract_tables_from_docx(rag_file)
        load_tables_and_append(tables=tables, docs=docs, file_name=rag_file)

    # DOC
    elif info.extension_matches("doc"):
        logger.info("RAG file " + rag_file + " was detected as a DOC file")
        #loader = Document(rag_file)
        loader = UnstructuredWordDocumentLoader(rag_file)
        load_pages_and_append(loader, docs)

    # ODT
    elif info.extension_matches("odt"):
        logger.info("RAG file " + rag_file + " was detected as an ODT file")

        # Use UnstructuredODTLoader first
        #loader = UnstructuredODTLoader(rag_file, mode="elements") # This doesn't work?
        loader = UnstructuredODTLoader(rag_file)
        load_pages_and_append(loader, docs)

        # Retrieve tables and append
        #tables = extract_tables_from_odt(rag_file)
        #load_tables_and_append(tables, docs)

    # PPTX, HTML
    elif info.extension_matches("pptx") or info.extension_matches("html"):
        logger.info("RAG file " + rag_file + " was detected as a PPTX/HTML file")
        loader = DoclingLoader(rag_file)
        load_pages_and_append(loader, docs)

    # CSV
    elif info.extension_matches("csv"):
        logger.info("RAG file " + rag_file + " was detected as a CSV file")
        loader = UnstructuredExcelLoader(rag_file)
        #load_pages_and_append(loader, docs)
        for page in loader.lazy_load():
            docs.append(page)
    # Other
    elif rag_file.lower().endswith(".txt"):
        logger.info("RAG file " + rag_file + " was detected as a TXT file")
        loader = TextLoader(rag_file)
        #load_pages_and_append(loader, docs)
        for page in loader.lazy_load():
            page.page_content = page.page_content.replace("　", " ")
            page.page_content = page.page_content.replace("\\u3000", " ")
            docs.append(page)

# Load page and append to docs 
def load_pages_and_append(loader, docs):
    doc = loader.load()
    doc = filter_complex_metadata(doc)
    for page in doc:
        page.page_content = page.page_content.replace("　", " ")
        docs.append(page)

# Load PDF with pdfplumber
def load_pdf_and_append(pdf_path, docs):
    # Get file name
    file_name = os.path.basename(pdf_path)

    # Initialize doc for texts in all pages which are not included in tables
    page_doc = Document(
                page_content="",
                metadata={"source": file_name, "page": 0}
                )

    # Open file with pdfplumber
    with pdfplumber.open(pdf_path) as pdf:
        p = 1
        t = 1
        for page in pdf.pages:
            #if p >= 2:
                #break

            # Extract tables from this page
            extracted_tables = page.extract_tables()

            for table in extracted_tables:
                table_str = ""

                if hasattr(table, 'rows'):
                    rows = table.rows
                else:
                    rows = table

                # Create string for each row
                for row in rows:
                    row_str = ""
                    if isinstance(row, list):
                        for cell in row:
                            if cell is None:
                                row_str += "\t"
                            else:
                                cell_str = str(cell).replace("　", " ")
                                #print(cell_str)
                                '''
                                if '\n' in cell_str:
                                    #cell_str = str(cell_str).replace("\n", "\t")
                                    row_str += "\n" + cell_str
                                else:
                                    row_str += "\t" + cell_str
                                '''
                                cell_str = str(cell_str).replace("\n", " ")
                                row_str += "\t" + cell_str
                    else:
                        row_str = str(row)

                    # Append row str to table str
                    table_str += "\n" + row_str

                # Append table str to docs       
                doc = Document(
                    page_content=table_str,
                    metadata={"source": file_name, "table_number": t}
                )
                t += 1
                docs.append(doc)

            # Now strings in tables are all appended to docs
            # Next we append text that are not included in any tables in page
            # Filter out all tables from this page first
            page = filter_tables_from_pdf(page)

            # Extract text 
            text = page.extract_text()

            # Append table str to docs       
            '''
            doc = Document(
                page_content=text,
                metadata={"source": file_name, "page": p}
            )
            docs.append(doc)
            '''
            page_doc.page_content += text 
            p += 1

    docs.append(page_doc)

# Filter out tables from PDF page
def filter_tables_from_pdf(page: pdfplumber.page.Page) -> pdfplumber.page.Page:
    if page.find_tables() != []:
        # Get the bounding boxes of the tables on the page.
        # Adapted from
        # https://github.com/jsvine/pdfplumber/issues/242#issuecomment-668448246
        bboxes = [table.bbox for table in page.find_tables()]
        bbox_not_within_bboxes = partial(not_within_bboxes, bboxes=bboxes)

        # Filter-out tables from page
        page = page.filter(bbox_not_within_bboxes)

    return page

def not_within_bboxes(obj, bboxes):
    """Check if the object is in any of the table's bbox."""

    def obj_in_bbox(_bbox):
        """Define objects in box.

        See https://github.com/jsvine/pdfplumber/blob/stable/pdfplumber/table.py#L404
        """
        v_mid = (obj["top"] + obj["bottom"]) / 2
        h_mid = (obj["x0"] + obj["x1"]) / 2
        x0, top, x1, bottom = _bbox
        return (h_mid >= x0) and (h_mid < x1) and (v_mid >= top) and (v_mid < bottom)

    return not any(obj_in_bbox(__bbox) for __bbox in bboxes)

# Extract tables from PDF
def extract_tables_from_pdf(pdf_path):
    # Get file name
    #file_name = os.path.basename(pdf_path)

    tables = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            extracted_tables = page.extract_tables()
            for table in extracted_tables:
                '''
                for row in table:
                    row_data = []
                    for cell in row:
                        row_data.append(str(cell))
                '''
                tables.append(table)  # Each table is a list of lists (rows)

    return tables

# Extract tables from DOCX file
def extract_tables_from_docx(file_path):
    # Get file name
    #file_name = os.path.basename(file_path)

    doc = Doc(file_path)
    tables = []
    
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text.strip() for cell in row.cells])
        tables.append(table_data)
    
    return tables

# Extract tables from ODT file
# This actually needs to be reviewed
# We do not use anyway
def extract_tables_from_odt(file_path):
    doc = load(file_path)
    tables = []

    for table in doc.getElementsByType(Table):
        table_data = []
        for row in table.getElementsByType(TableRow):
            row_data = []
            for cell in row.getElementsByType(TableCell):
                ''' This complains p.firstChild has no 'data' attr
                for p in cell.getElementsByType(P):
                    print(dir(p))
                    print(p)
                    q = p.firstChild
                    print(dir(q))
                    print(q)
                cell_text = " ".join(p.firstChild.data for p in cell.getElementsByType(P) if p.firstChild)
                '''

                # Hiroshi's version
                cell_text = ""
                for p in cell.getElementsByType(P):
                    if p is not None:
                        cell_text += str(p).replace("　", " ") + " "

                '''
                print("cell: " + str(cell))
                print("cell_text: " + cell_text + "\n")
                '''

                # Add cells to row
                row_data.append(cell_text)
            table_data.append(row_data)
        tables.append(table_data)
    
    return tables

# Append content of tables returned by extract_tables_from_XXX() to docs
def load_tables_and_append(
        tables,
        docs,
        file_name
        ):
    for t, table in enumerate(tables):
        table_str = ""

        # Check table has attribute 'rows'
        rows = None
        if hasattr(table, 'rows'):
            rows = table.rows
        else:
            rows = table

        # If row is list, format row as string
        for row in rows:
            row_str = ""
            if isinstance(row, list):
                for cell in row:
                    if cell is None:
                        row_str += "\t"
                    else:
                        cell_str = str(cell).replace("　", " ")
                        row_str += "\t" + cell_str
            else:
                row_str = str(row)

        table_str += "\n" + row_str

        #table_str = "\n".join(["\t".join(row) for row in table])  # Format as tab-separated text
        doc = Document(
            page_content=table_str,
            metadata={"source": file_name, "table_number": t+1}
        )
        docs.append(doc)

# Load websites
def load_websites(
        urls: list[str],
        chunk_size = 1024,
        chunk_overlap = 120,
        max_links = 20
        ):
    # Access websites and get docs
    '''
    loader = WebBaseLoader(urls[0])
    docs = loader.load()
    '''
    try:
        loader = MyWebLoader(urls=urls, max_links=max_links)
    except Exception as e:
        raise Exception(e)
        
    docs, dummy, combined_ctx = loader.load()
    #logger.debug(docs)
    #exit()

    # Initialize text splitter
    text_splitter = RecursiveCharacterTextSplitter(
                        separators=[
                            ".\n\n",
                            "\n\n",
                            "\r\n\r\n",
                            "。",
                        ],
                        keep_separator='end',
                        chunk_size=chunk_size,
                        chunk_overlap=chunk_overlap,
                        add_start_index=True,
                        ) 

    # Split docs
    docs = text_splitter.split_documents(docs)

    #return docs, combined_ctx
    return CtxDocs(docs=docs, combined_ctx=combined_ctx)

# Load raw text/html data
def load_raw_data(
        data: list[str],
        chunk_size = 1024,
        chunk_overlap = 120
        ):

    docs = []
    page = 1

    for d in data:
        text = None
        soup = BeautifulSoup(d, 'html.parser')
        if bool(soup.find()):
            # HTML
            text = soup.get_text()
        else:
            # Text
            text = d
        
        # Create document and append
        doc = Document(
                page_content=text,
                metadata={"source": "copied_data", "page": page},
                )
        page += 1
        docs.append(doc)

    # Initialize text splitter
    text_splitter = RecursiveCharacterTextSplitter(
                        separators=[
                            ".\n\n",
                            "\n\n",
                            "\r\n\r\n",
                            "。",
                        ],
                        keep_separator='end',
                        chunk_size=chunk_size,
                        chunk_overlap=chunk_overlap,
                        add_start_index=True,
                        ) 

    # Split docs
    docs = text_splitter.split_documents(docs)

    return docs

# Extract URLs from sitemap XML
def extract_urls_from_sitemap(sitemap):
    """
    Extract all URLs from a sitemap XML string.

    Args:
        sitemap_string (str): The sitemap XML string.

    Returns:
        A list of URLs extracted from the sitemap.
    """
    # Parse the XML from the string
    parser = ET.XMLParser(encoding="utf-8")
    root = ET.fromstring(sitemap, parser=parser)

    # Define the namespace for the sitemap XML
    namespace = {"ns": "http://www.sitemaps.org/schemas/sitemap/0.9"}

    # Find all <loc> elements under the <url> elements
    urls = [
        url.find("ns:loc", namespace).text for url in root.findall("ns:url", namespace)
    ]

    # Return the list of URLs
    return urls

# Anonymize doc source in docs
def anonymize_source_in_docs(docs):
    for doc in docs:
        doc.metadata['source'] = \
            ''.join(random.choices(string.ascii_letters + string.digits, k=8))

# Replace text in docs
def replace_text_in_docs(docs, old_text, new_text):
    for doc in docs:
        doc.page_content = doc.page_content.replace(old_text, new_text)

# Replace human name in docs
def replace_name_in_docs(
                 docs: list[Document],
                 first_name: str,
                 last_name: str,
                 new_first_name: str = "AnonymousFirstName",
                 new_last_name: str = "AnonymousLastName",
                 ):
    # Initialize
    first_name_len = len(first_name)
    last_name_len = len(last_name)

    # Check last_name_len and first_name_len
    if last_name_len == 0 or first_name_len == 0:
        # Name not complete, we use either of them 
        name = first_name + last_name

        # Replace name
        replace_text_in_docs(docs, name, new_last_name + new_first_name)

    else:
        for doc in docs:
            start_idx = 0
            idx = 0
            page_content = doc.page_content
            while True:
                page_content_len = len(page_content)
                if start_idx >= page_content_len - 1:
                    break

                # Find first_name
                first_name_idx = page_content.find(first_name, start_idx)
                #print("first_name_idx: " + str(first_name_idx))
                if first_name_idx < 0:
                    break
                elif first_name_idx == 0:
                    page_content = new_first_name + page_content[1:]
                    start_idx = 1
                    continue
    
                if first_name_idx > last_name_len + 3:
                    idx_s = first_name_idx - last_name_len - 3
                else:
                    idx_s = 0
    
                if first_name_idx < page_content_len - first_name_len - 3:
                    idx_e = first_name_idx + first_name_len + 3
                else:
                    idx_e = page_content_len
    
                # Find nearest last_name
                last_name_idx = page_content.find(last_name, idx_s, idx_e)
                #print("last_name_idx: " + str(last_name_idx))
                if last_name_idx < 0:
                    start_idx = first_name_idx + 1
                    break

                # Last name found near first name
                # Replace them
                if first_name_idx < last_name_idx:
                    page_content = \
                        page_content[:first_name_idx] + new_first_name + \
                        page_content[first_name_idx+first_name_len:last_name_idx] + \
                        new_last_name + \
                        page_content[last_name_idx+last_name_len:]
                else:
                    page_content = \
                        page_content[:last_name_idx] + new_last_name + \
                        page_content[last_name_idx+last_name_len:first_name_idx] + \
                        new_first_name + \
                        page_content[first_name_idx+first_name_len:]
    
                doc.page_content = page_content
    
# Main
if __name__ == "__main__":
    # Load files
    docs, _ = load_files(files=sys.argv[1:])
    print(docs)

    # Replace name
    replace_name_in_docs(docs, last_name="イプシロン田中", first_name="")
    #replace_name_in_docs(docs, last_name="西田", first_name="博史")
    #replace_name_in_docs(docs, last_name="ニシダ", first_name="ヒロシ")
    print("\n")
    print(docs)

    # Replace others
    replace_text_in_docs(docs, "nishidafmly@aol.com", "anonymouse_email")
    replace_text_in_docs(docs, "1683 Ammon St NW, Salem", "anonymouse_address")
    replace_text_in_docs(docs, "503-779-3167", "anonymouse_phone")
    replace_text_in_docs(docs, "asusa.com", "anonymouse_url")
    print("\n")
    print(docs)
